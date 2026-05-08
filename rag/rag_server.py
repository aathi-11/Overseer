from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
import chromadb
import httpx
import logging
import csv
import json
import uuid
import io
import re
import base64
import asyncio

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

chroma = chromadb.PersistentClient(path="./chroma_db")
collection = chroma.get_or_create_collection("overseer_knowledge", metadata={"hnsw:space": "cosine"})

OLLAMA_BASE_URL   = "http://localhost:11434"
OLLAMA_EMBED_URL  = f"{OLLAMA_BASE_URL}/api/embeddings"
OLLAMA_GEN_URL    = f"{OLLAMA_BASE_URL}/api/generate"
EMBED_MODEL       = "nomic-embed-text"
VISION_MODEL      = "gemma3:4b"

# ── Cosine distance threshold ─────────────────────────────────────────────────
# Results with cosine distance > this value are considered irrelevant.
# 0.6 gives good recall for dense financial/technical docs; tighten to 0.4 if
# you start getting unrelated chunks mixed in.
COSINE_DISTANCE_THRESHOLD = 0.6


async def get_embedding(text: str):
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            res = await client.post(OLLAMA_EMBED_URL, json={
                "model": EMBED_MODEL,
                "prompt": text
            })
            res.raise_for_status()
            data = res.json()
            return data.get("embedding")
    except Exception as e:
        logger.warning("Embedding failed: %s", e)
        return None


class QueryRequest(BaseModel):
    query: str
    n_results: int = 3
    # Optional metadata filter for per-agent RAG queries
    # e.g. {"type": "pattern"} or {"type": "test"} or {"type": "requirement"}
    where: Optional[dict] = None
    # Optional doc_id to scope queries to a single ingested document
    doc_id: Optional[str] = None

class StoreRequest(BaseModel):
    id: str
    content: str
    metadata: dict = {}

@app.post("/query")
async def query_knowledge(req: QueryRequest):
    try:
        count = collection.count()
        if count == 0:
            return {"chunks": [], "ids": [], "distances": [], "found": 0}
        actual_n = min(req.n_results, count)
        embedding = await get_embedding(req.query)
        if not embedding:
            return {"chunks": [], "ids": [], "distances": [], "found": 0, "error": "embedding_failed"}

        query_kwargs = {
            "query_embeddings": [embedding],
            "n_results": actual_n,
            "include": ["documents", "distances", "metadatas"],
        }

        # Build the where clause — doc_id takes priority (most specific filter)
        where_clause = {}
        if req.doc_id:
            # Scope to exactly this uploaded document — prevents cross-session contamination
            where_clause = {"doc_id": {"$eq": req.doc_id}}
        elif req.where and isinstance(req.where, dict) and len(req.where) > 0:
            where_clause = req.where

        if where_clause:
            query_kwargs["where"] = where_clause

        results = collection.query(**query_kwargs)
        docs      = results["documents"][0] if results["documents"] else []
        ids       = results["ids"][0] if results["ids"] else []
        distances = results["distances"][0] if results.get("distances") else []

        logger.debug("RAG query doc_id=%s distances: %s", req.doc_id, distances)

        # ── Cosine distance filter ────────────────────────────────────────────
        # Drop any chunk whose distance is above the threshold — irrelevant results
        filtered_docs = []
        filtered_ids  = []
        for doc, did, dist in zip(docs, ids, distances):
            if dist <= COSINE_DISTANCE_THRESHOLD:
                filtered_docs.append(doc)
                filtered_ids.append(did)
            else:
                logger.debug("Filtered out chunk id=%s distance=%.3f", did, dist)

        return {"chunks": filtered_docs, "ids": filtered_ids, "found": len(filtered_docs)}
    except Exception as e:
        logger.error("Query error: %s", e)
        return {"chunks": [], "ids": [], "found": 0, "error": str(e)}

@app.post("/store")
async def store_knowledge(req: StoreRequest):
    try:
        embedding = await get_embedding(req.content)
        if not embedding:
            return {"stored": False, "error": "embedding_failed"}
        collection.upsert(
            ids=[req.id],
            documents=[req.content],
            embeddings=[embedding],
            metadatas=[req.metadata]
        )
        logger.info("Stored document id=%s type=%s", req.id, req.metadata.get("type", "unknown"))
        return {"stored": True, "id": req.id}
    except Exception as e:
        logger.error("Store error: %s", e)
        return {"stored": False, "error": str(e)}

@app.on_event("startup")
async def on_startup():
    try:
        count = collection.count()
        logger.info("RAG server ready. Collection size: %s documents.", count)
        if count == 0:
            logger.warning("RAG collection is empty — run seed_knowledge.py to populate it.")
    except Exception as e:
        logger.warning("RAG server startup check failed: %s", e)

@app.get("/health")
async def health_check():
    count = collection.count()
    # Keep both keys: 'count' for ragClient.js compatibility, 'db_count' for humans
    return {"status": "ok", "count": count, "db_count": count}


# ─────────────────────────────────────────────────────────────────────────────
# PDF VISION EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

async def extract_page_with_vision(page_image_bytes: bytes, page_num: int) -> str:
    """
    Send a rendered PDF page PNG to gemma3:4b via Ollama /api/generate.
    Returns the extracted text content for that page.
    """
    b64_image = base64.b64encode(page_image_bytes).decode("utf-8")
    payload = {
        "model": VISION_MODEL,
        "prompt": (
            f"This is page {page_num} of a document. "
            "Extract all text content visible in this image. "
            "Return only the raw text — no commentary, no formatting, just the content."
        ),
        "images": [b64_image],
        "stream": False,
    }
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            res = await client.post(OLLAMA_GEN_URL, json=payload)
            res.raise_for_status()
            data = res.json()
            return data.get("response", "").strip()
    except Exception as e:
        logger.warning("Vision extraction failed for page %d: %s", page_num, e)
        return ""


async def extract_pdf_chunks_smart(content: bytes, chunk_size: int = 500, progress_cb=None) -> list:
    """
    Hybrid PDF chunker:
    - Text-heavy pages (>150 chars) → PyMuPDF text extraction (fast)
    - Image/design pages (<= 150 chars) → gemma3:4b vision via Ollama
    Each chunk is prefixed with [Page N] for citation.

    progress_cb: optional async callable(page_num, total_pages) called after each page.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF (fitz) not installed — falling back to pdfplumber")
        return extract_pdf_chunks_pdfplumber(content, chunk_size)

    chunks = []
    doc = fitz.open(stream=content, filetype="pdf")
    total_pages = len(doc)
    logger.info("PDF has %d pages — using smart hybrid extraction", total_pages)

    for page_num in range(total_pages):
        page = doc[page_num]
        page_label = page_num + 1  # 1-indexed for display

        # ── Try text extraction first ─────────────────────────────────────────
        text = page.get_text("text").strip()

        if len(text) > 150:
            # Enough text — chunk by paragraph
            for para in text.split("\n\n"):
                para = para.strip()
                if len(para) > 50:
                    chunk = f"[Page {page_label}] {para[:chunk_size]}"
                    chunks.append(chunk)
            logger.debug("Page %d: text extraction (%d chars)", page_label, len(text))
        else:
            # Low text — try vision
            logger.info("Page %d: low text (%d chars) — using vision model", page_label, len(text))
            try:
                # Render page to PNG at 150 DPI
                mat = fitz.Matrix(150 / 72, 150 / 72)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")

                vision_text = await extract_page_with_vision(img_bytes, page_label)
                if vision_text:
                    for para in vision_text.split("\n\n"):
                        para = para.strip()
                        if len(para) > 50:
                            chunk = f"[Page {page_label}] {para[:chunk_size]}"
                            chunks.append(chunk)
                    logger.debug("Page %d: vision extracted %d chars", page_label, len(vision_text))
                else:
                    logger.debug("Page %d: vision returned empty — skipping", page_label)
            except Exception as ve:
                logger.warning("Page %d: vision failed, skipping: %s", page_label, ve)

        # ── Emit progress ─────────────────────────────────────────────────────
        if progress_cb:
            await progress_cb(page_label, total_pages)

    doc.close()
    return chunks


def extract_pdf_chunks_pdfplumber(content: bytes, chunk_size: int = 500) -> list:
    """Fallback: uses pdfplumber if PyMuPDF is not available."""
    import pdfplumber
    chunks = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            page_label = i + 1
            for para in text.split("\n\n"):
                para = para.strip()
                if len(para) > 50:
                    chunks.append(f"[Page {page_label}] {para[:chunk_size]}")
    return chunks


def extract_excel_chunks(content: bytes, rows_per_chunk: int = 5) -> list:
    chunks = []
    try:
        import openpyxl
        import io as _io
        wb = openpyxl.load_workbook(_io.BytesIO(content), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers = []
            batch_lines = []
            rows_seen = 0
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(v) if v is not None else "" for v in row]
                if not any(row_vals):
                    continue
                if not headers:
                    headers = row_vals
                    continue
                line = " | ".join(
                    f"{headers[i] if i < len(headers) else i}: {v}"
                    for i, v in enumerate(row_vals)
                )
                batch_lines.append(line)
                rows_seen += 1
                if rows_seen % rows_per_chunk == 0:
                    chunks.append(f"[Sheet: {sheet_name}]\n" + "\n".join(batch_lines))
                    batch_lines = []
            if batch_lines:
                chunks.append(f"[Sheet: {sheet_name}]\n" + "\n".join(batch_lines))
        wb.close()
    except Exception as e:
        logger.warning("Excel extraction failed: %s", e)
    return chunks


def extract_docx_chunks(content: bytes, chunk_size: int = 500) -> list:
    chunks = []
    try:
        import docx
        import io as _io
        doc = docx.Document(_io.BytesIO(content))
        full_text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            if rows:
                full_text += "\n\n[Table]\n" + "\n".join(rows)
        for i in range(0, len(full_text), chunk_size):
            chunk = full_text[i:i + chunk_size].strip()
            if chunk:
                chunks.append(chunk)
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
    return chunks


async def extract_image_chunks(content: bytes, filename: str) -> list:
    b64_image = base64.b64encode(content).decode("utf-8")
    payload = {
        "model": VISION_MODEL,
        "prompt": (
            f"This is an uploaded image named '{filename}'. "
            "1. Describe everything you see in detail.\n"
            "2. Extract any visible text or data verbatim.\n"
            "Separate the two sections with '---'."
        ),
        "images": [b64_image],
        "stream": False,
    }
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            res = await client.post(OLLAMA_GEN_URL, json=payload)
            res.raise_for_status()
            text = res.json().get("response", "").strip()
            if text:
                return [text[i:i+500] for i in range(0, len(text), 500) if text[i:i+500].strip()]
    except Exception as e:
        logger.warning("Image extraction failed for %s: %s", filename, e)
    return [f"[Image: {filename}] Could not extract content."]


# ─────────────────────────────────────────────────────────────────────────────
# INGEST ENDPOINT — with SSE progress streaming for PDFs
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/ingest")
async def ingest_document(
    file: UploadFile = File(...),
    doc_type: str = Form("document"),
    filename: str = Form("upload")
):
    """
    Ingest a document into ChromaDB.

    For PDFs: streams SSE progress events as each page is processed so the
    frontend can display "Processing page N of total…" messages.

    Returns JSON for non-PDF files. For PDFs it also returns JSON (progress
    events are emitted before the final response).
    """
    try:
        content = await file.read()
        effective_filename = filename if filename != "upload" else (file.filename or "upload")
        ext = effective_filename.rsplit(".", 1)[-1].lower()

        logger.info("Ingest: filename=%s ext=%s size=%d bytes", effective_filename, ext, len(content))

        doc_id = str(uuid.uuid4())[:8]

        if ext == "pdf":
            # ── Stream SSE progress for PDFs ──────────────────────────────────
            async def pdf_stream():
                progress_messages = []

                async def on_page_progress(page_num, total):
                    msg = f"Processing page {page_num} of {total}…"
                    progress_messages.append(msg)
                    yield f"data: {json.dumps({'type': 'progress', 'message': msg})}\n\n"

                # We need to collect chunks AND stream at the same time.
                # We'll use a generator approach: wrap extract_pdf_chunks_smart
                # so it yields SSE events as a side effect.
                chunks_holder = []

                async def extract_with_progress():
                    async def _cb(page_num, total):
                        progress_messages.append(f"Processing page {page_num} of {total}…")

                    result = await extract_pdf_chunks_smart(content, progress_cb=_cb)
                    chunks_holder.extend(result)

                # Run extraction in background while also streaming progress
                # We use a queue-based approach to stream in real time
                progress_queue: asyncio.Queue = asyncio.Queue()

                async def _progress_cb(page_num, total):
                    await progress_queue.put((page_num, total))

                extraction_task = asyncio.create_task(
                    extract_pdf_chunks_smart(content, progress_cb=_progress_cb)
                )

                # Stream progress while extraction is running
                while not extraction_task.done():
                    try:
                        page_num, total = await asyncio.wait_for(
                            progress_queue.get(), timeout=1.0
                        )
                        msg = f"Processing page {page_num} of {total}…"
                        yield f"data: {json.dumps({'type': 'progress', 'message': msg})}\n\n"
                    except asyncio.TimeoutError:
                        continue

                # Drain remaining progress events
                while not progress_queue.empty():
                    page_num, total = progress_queue.get_nowait()
                    msg = f"Processing page {page_num} of {total}…"
                    yield f"data: {json.dumps({'type': 'progress', 'message': msg})}\n\n"

                chunks = await extraction_task
                logger.info("PDF extracted %d chunks", len(chunks))

                if not chunks:
                    yield f"data: {json.dumps({'type': 'done', 'doc_id': None, 'chunk_count': 0, 'total_chunks': 0, 'warning': 'No extractable text found in PDF.'})}\n\n"
                    return

                # Embed + store chunks
                stored = 0
                total_chunks_to_embed = len(chunks)
                for i, chunk in enumerate(chunks):
                    # Periodically yield progress so the frontend doesn't stall
                    if i % 10 == 0 or i == total_chunks_to_embed - 1:
                        msg = f"Embedding chunk {i + 1} of {total_chunks_to_embed}…"
                        yield f"data: {json.dumps({'type': 'progress', 'message': msg})}\n\n"

                    embedding = await get_embedding(chunk)
                    if not embedding:
                        logger.warning("Skipping chunk %d — embedding failed", i)
                        continue
                    try:
                        collection.add(
                            documents=[chunk],
                            embeddings=[embedding],
                            ids=[f"{doc_id}-{i}"],
                            metadatas=[{
                                "doc_id": doc_id,
                                "doc_type": doc_type,
                                "filename": effective_filename,
                                "chunk_index": i
                            }]
                        )
                        stored += 1
                    except Exception as chunk_err:
                        logger.warning("Failed to store chunk %d: %s", i, chunk_err)

                logger.info("PDF ingest complete doc_id=%s stored=%d", doc_id, stored)
                yield f"data: {json.dumps({'type': 'done', 'doc_id': doc_id, 'chunk_count': stored, 'total_chunks': len(chunks)})}\n\n"

            return StreamingResponse(pdf_stream(), media_type="text/event-stream")

        else:
            # ── Non-PDF: extract + store synchronously, return JSON ────────────
            if ext == "csv":
                chunks = extract_csv_chunks(content.decode("utf-8", errors="replace"))
            elif ext == "json":
                chunks = extract_json_chunks(content.decode("utf-8", errors="replace"))
            elif ext in ("xlsx", "xls"):
                chunks = extract_excel_chunks(content)
            elif ext == "docx":
                chunks = extract_docx_chunks(content)
            elif ext in ("png", "jpg", "jpeg", "webp", "gif"):
                chunks = await extract_image_chunks(content, effective_filename)
            else:
                text = content.decode("utf-8", errors="replace")
                chunks = [text[i:i+500] for i in range(0, len(text), 500) if text[i:i+500].strip()]

            logger.info("Extracted %d chunk(s) from %s", len(chunks), effective_filename)

            if not chunks:
                return {"doc_id": None, "chunk_count": 0, "warning": "No extractable text found in file."}

            stored = 0
            for i, chunk in enumerate(chunks):
                embedding = await get_embedding(chunk)
                if not embedding:
                    logger.warning("Skipping chunk %d — embedding failed", i)
                    continue
                try:
                    collection.add(
                        documents=[chunk],
                        embeddings=[embedding],
                        ids=[f"{doc_id}-{i}"],
                        metadatas=[{
                            "doc_id": doc_id,
                            "doc_type": doc_type,
                            "filename": effective_filename,
                            "chunk_index": i
                        }]
                    )
                    stored += 1
                except Exception as chunk_err:
                    logger.warning("Failed to store chunk %d: %s", i, chunk_err)

            logger.info("Ingested doc_id=%s type=%s file=%s stored=%d total_chunks=%d",
                        doc_id, doc_type, effective_filename, stored, len(chunks))
            return {"doc_id": doc_id, "chunk_count": stored, "total_chunks": len(chunks)}

    except Exception as e:
        logger.error("Ingest failed: %s", e, exc_info=True)
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────────────────────────────────────
# CSV CHUNKING
# ─────────────────────────────────────────────────────────────────────────────

def repair_wide_headers(headers: list) -> list:
    """Replace empty / ######## headers by incrementing from the last known year.
    Excel truncates column widths visually, but the CSV export should have real dates.
    This repairs any that slipped through as empty strings.
    """
    repaired = []
    last_year = None
    for h in headers:
        h = str(h).strip()
        if not h or h == "########":
            if last_year:
                last_year += 1
                repaired.append(f"Year-{last_year}")
            else:
                repaired.append("Unknown")
        else:
            repaired.append(h)
            match = re.search(r'(\d{4})', h)
            if match:
                last_year = int(match.group(1))
    return repaired


def extract_csv_chunks(content: str, rows_per_chunk: int = 5) -> list:
    """Smart CSV chunker that handles two layouts:

    WIDE format (financial balance sheets):
      Row 0  : headers = ["", "2011", "2012", "2013", ...]
      Row 1+ : data    = ["Borrowings", 0, 328, 620, ...]
      → Transpose: one chunk per metric row
         e.g. "Borrowings: 2011=0, 2012=328, 2013=620"

    TALL format (normal transaction logs):
      Row 0  : headers = ["Date", "Company", "Amount"]
      Row 1+ : data    = ["2024-01", "Acme", "50000"]
      → Chunk rows_per_chunk rows at a time
    """
    chunks = []
    try:
        reader = csv.reader(io.StringIO(content))
        rows = [r for r in reader if any(c.strip() for c in r)]  # skip blank rows

        if not rows:
            logger.warning("CSV has no data rows — returning raw content")
            return [content[:2000]]

        raw_headers = rows[0]
        data_rows   = rows[1:]

        # Repair ######### / empty header cells
        headers = repair_wide_headers(raw_headers)

        print(f"[CSV] {len(data_rows)} data rows, {len(headers)} columns: {headers[:6]}...")
        logger.info("CSV: %d data rows, %d columns", len(data_rows), len(headers))

        # ── Detect wide format ────────────────────────────────────────────────
        # Wide: first header is blank/label, remaining headers are years/dates
        is_wide = (
            len(headers) > 2
            and (not headers[0].strip() or not re.search(r'\d{4}', headers[0]))
            and sum(1 for h in headers[1:] if re.search(r'\d{4}', str(h))) >= 2
        )

        if is_wide:
            logger.info("Detected WIDE format — transposing to metric-per-chunk")
            print("[CSV] WIDE format detected — one chunk per metric row")
            for row in data_rows:
                if not row or not row[0].strip():
                    continue
                metric = row[0].strip()
                values = []
                for i, val in enumerate(row[1:], start=1):
                    if i < len(headers):
                        year_label = headers[i].strip()
                        val_clean  = val.strip()
                        if year_label and val_clean:
                            values.append(f"{year_label}={val_clean}")
                if values:
                    chunk = f"{metric}: {', '.join(values)}"
                    chunks.append(chunk)
                    print(f"  → {chunk[:80]}")
        else:
            # ── Tall format: chunk by rows_per_chunk ─────────────────────────
            logger.info("Detected TALL format — chunking %d rows at a time", rows_per_chunk)
            for i in range(0, len(data_rows), rows_per_chunk):
                batch = data_rows[i:i + rows_per_chunk]
                lines = [
                    " | ".join(
                        f"{headers[j] if j < len(headers) else j}: {v}"
                        for j, v in enumerate(row)
                    )
                    for row in batch if row
                ]
                if lines:
                    chunks.append("\n".join(lines))

        print(f"[CSV] produced {len(chunks)} chunks")
        logger.info("CSV produced %d chunks", len(chunks))

    except Exception as e:
        logger.warning("CSV chunking failed, raw fallback: %s", e, exc_info=True)
        chunks = [content[i:i+500] for i in range(0, len(content), 500) if content[i:i+500].strip()]

    return chunks


def extract_json_chunks(content: str):
    data = json.loads(content)
    # Flatten nested JSON into readable chunks
    if isinstance(data, list):
        chunks = [json.dumps(item, indent=2) for item in data]
    elif isinstance(data, dict):
        chunks = [f"{k}: {json.dumps(v)}" for k, v in data.items()]
    else:
        chunks = [content]
    return chunks
